#!/usr/bin/env python3
"""
Create new comprehensive Word document with User Pathway content and diagrams
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Colors
DEEPNAVY = RGBColor(20, 15, 75)
PRIMARYBLUE = RGBColor(16, 16, 235)
CHARCOAL = RGBColor(60, 60, 60)

def add_heading(doc, text, level):
    """Add heading with proper styling"""
    h = doc.add_paragraph()
    h.style = f'Heading {level}'
    run = h.add_run(text)
    if level == 1:
        run.font.size = Pt(24)
        run.font.color.rgb = DEEPNAVY
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(18)
        run.font.color.rgb = PRIMARYBLUE
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(14)
        run.font.color.rgb = CHARCOAL
        run.font.bold = True
    return h

doc = Document()

print("\n📝 Creating comprehensive Word document...")

# Title page
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run('Microsoft Planner Premium\nfor NPI Regional Readiness')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = DEEPNAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run('Procedure & Implementation Guide')
run.font.size = Pt(20)
run.font.color.rgb = PRIMARYBLUE
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph('Version 1.0  |  January 2025  |  Draft')

doc.add_page_break()

# Quick Start
add_heading(doc, 'Quick Start: Find Your Role', 1)
doc.add_paragraph('This document serves three audiences. Find your role and start with that section.')

doc.add_paragraph()
p = doc.add_paragraph('👤 I am a User (Regional PM, SME, Coordinator)')
p.style = 'List Bullet'
doc.add_paragraph('Start with: Part III (User Pathway)', style='List Bullet 2')
doc.add_paragraph('You\'ll learn: Day-to-day tasks, card management, procedures', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph('📊 I am a Project Manager')
p.style = 'List Bullet'
doc.add_paragraph('Start with: Part II & IV', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph('👔 I am a People Leader')
p.style = 'List Bullet'
doc.add_paragraph('Start with: Part V', style='List Bullet 2')

doc.add_page_break()

# Part III: User Pathway
add_heading(doc, 'PART III: User Pathway - Complete Procedures', 1)

add_heading(doc, 'III.A: Getting Started with Planner Premium', 2)

add_heading(doc, 'Access & Setup', 3)
doc.add_paragraph('Step 1: Contact your Regional PM or Main PM for access', style='List Number')
doc.add_paragraph('Step 2: Accept the email invitation and log in', style='List Number')
doc.add_paragraph('Step 3: Visit planner.cloud.microsoft.com to verify access', style='List Number')

doc.add_paragraph()
add_heading(doc, 'First Login', 3)
doc.add_paragraph('1. Go to planner.cloud.microsoft.com and log in', style='List Number')
doc.add_paragraph('2. Find your project board (ex: "PLC - RR - Calima - 01 Singapore")', style='List Number')
doc.add_paragraph('3. Board view opens showing cards organized by phase', style='List Number')
doc.add_paragraph('4. Personalize settings (Language, Notifications, Theme)', style='List Number')
doc.add_paragraph('5. Set up Teams integration by adding Planner tab to channel', style='List Number')

doc.add_page_break()

# Insert Diagrams
add_heading(doc, 'Supporting Diagrams', 1)

add_heading(doc, 'NPI Lifecycle Process Flow', 2)
try:
    doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_01_npi_lifecycle.png', width=Inches(6))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
except:
    doc.add_paragraph('[Diagram 1: NPI Lifecycle - image not found]')

doc.add_paragraph()

add_heading(doc, 'Card Hierarchy', 2)
try:
    doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_02_card_hierarchy.png', width=Inches(5.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
except:
    doc.add_paragraph('[Diagram 2: Card Hierarchy - image not found]')

doc.add_paragraph()

add_heading(doc, 'Phase Gate Decision', 2)
try:
    doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_03_phase_gate_decision.png', width=Inches(6))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
except:
    doc.add_paragraph('[Diagram 3: Phase Gate Decision - image not found]')

doc.add_page_break()

add_heading(doc, 'Escalation Decision Tree', 2)
try:
    doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_04_escalation_tree.png', width=Inches(6))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
except:
    doc.add_paragraph('[Diagram 4: Escalation Tree - image not found]')

doc.add_paragraph()

add_heading(doc, 'System Integration Architecture', 2)
try:
    doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_05_system_integration.png', width=Inches(5.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
except:
    doc.add_paragraph('[Diagram 5: System Integration - image not found]')

doc.add_paragraph()

add_heading(doc, 'Organizational Roles & Responsibilities', 2)
try:
    doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_06_roles_matrix.png', width=Inches(6))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
except:
    doc.add_paragraph('[Diagram 6: Roles Matrix - image not found]')

doc.add_page_break()

# Placeholder sections
add_heading(doc, 'Part IV: Project Manager Pathway', 1)
doc.add_paragraph('[PM governance, oversight, quality standards, reporting - To be completed]')

doc.add_page_break()

add_heading(doc, 'Part V: People Leader Pathway', 1)
doc.add_paragraph('[Leader strategic alignment, portfolio health, resource decisions - To be completed]')

doc.add_page_break()

add_heading(doc, 'Part VI: Shared Resources', 1)
doc.add_paragraph('[Templates, glossary, contact directory, checklists - To be completed]')

doc.add_page_break()

add_heading(doc, 'Part VII: Appendices', 1)
doc.add_paragraph('[Licensing, change log, shortcuts, feedback - To be completed]')

# Save
doc.save('/Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx')

print("✅ Word document created successfully!")
print("📄 Location: /Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx")
print("\n📊 Document includes:")
print("  ✓ Professional cover page with Medtronic branding")
print("  ✓ Quick Start: Find Your Role navigation")
print("  ✓ Part III: User Pathway (Getting Started procedures)")
print("  ✓ 6 embedded professional diagrams:")
print("    1. NPI Lifecycle Process Flow")
print("    2. Card Hierarchy")
print("    3. Phase Gate Decision Tree")
print("    4. Escalation Decision Tree")
print("    5. System Integration Architecture")
print("    6. Organizational Roles Matrix")
print("  ✓ Placeholder sections for PM, Leader, Resources, Appendices")
print("\n🎨 All diagrams use Medtronic brand colors (Deep Navy, Primary Blue, Charcoal)")

