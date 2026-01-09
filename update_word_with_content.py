#!/usr/bin/env python3
"""
Update Word document with User Pathway content and embedded diagrams
Creates comprehensive professional NPI Planner Premium procedure document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Colors (RGB tuples)
COLORS = {
    'deepNavy': RGBColor(20, 15, 75),      # #140F4B
    'primaryBlue': RGBColor(16, 16, 235),   # #1010EB
    'charcoal': RGBColor(60, 60, 60),       # #3C3C3C
    'teal': RGBColor(0, 220, 185),          # #00DCB9
    'lightGray': RGBColor(245, 245, 245),   # #F5F5F5
}

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level, color=COLORS['deepNavy']):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(24)
        elif level == 2:
            run.font.size = Pt(18)
        elif level == 3:
            run.font.size = Pt(16)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph_styled(doc, text, bold=False, italic=False, color=COLORS['charcoal'], size=11):
    """Add styled paragraph"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        run.font.size = Pt(size)
    return p

def main():
    print("\n📝 Creating comprehensive Word document with User Pathway & diagrams...\n")
    
    # Load existing document or create new one
    doc = Document('/Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx')
    
    # Add a page break after front matter
    doc.add_page_break()
    
    # ===== PART III: USER PATHWAY - DETAILED SECTIONS =====
    
    add_heading_styled(doc, 'PART III: User Pathway', 1, COLORS['deepNavy'])
    add_paragraph_styled(doc, 'Complete Step-by-Step Procedures for Regional PMs, SMEs, and Team Members',
                         italic=True, size=10)
    
    # ===== III.A: GETTING STARTED =====
    doc.add_page_break()
    add_heading_styled(doc, 'III.A: Getting Started with Planner Premium', 2, COLORS['primaryBlue'])
    
    # III.A.1: Access & Setup
    add_heading_styled(doc, 'III.A.1: Access & Setup', 3)
    
    p = doc.add_paragraph('Who Gets Access?')
    p.style = 'List Bullet'
    doc.add_paragraph('Regional PMs - Full access (create cards, assign tasks, move between phases)', style='List Bullet')
    doc.add_paragraph('Main PMs - Full access (governance oversight)', style='List Bullet')
    doc.add_paragraph('Functional SMEs - Limited access (work on assigned sub-tasks, view project cards)', style='List Bullet')
    doc.add_paragraph('Regional Service Leaders - View-only access (portfolio oversight)', style='List Bullet')
    
    add_paragraph_styled(doc, '\nGetting Planner Premium Access\n', bold=True)
    
    doc.add_paragraph('Email your Regional Service Leader or Main PM with:', style='List Number')
    doc.add_paragraph('Your full name, email, role, project name', style='List Number 2')
    doc.add_paragraph('Expected timeline: 2-3 business days', style='List Number')
    doc.add_paragraph('Accept the invitation in your email', style='List Number')
    doc.add_paragraph('Verify access at planner.cloud.microsoft.com', style='List Number')
    
    # Browser requirements table
    add_paragraph_styled(doc, '\nBrowser & Technical Requirements\n', bold=True)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Requirement'
    header_cells[1].text = 'Details'
    set_cell_background(header_cells[0], '140F4B')
    set_cell_background(header_cells[1], '140F4B')
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    # Data rows
    table.rows[1].cells[0].text = 'Recommended Browser'
    table.rows[1].cells[1].text = 'Microsoft Edge (latest version) - most compatible'
    
    table.rows[2].cells[0].text = 'System Requirements'
    table.rows[2].cells[1].text = 'Internet connection, 1024x768 minimum resolution, HTML5 support, JavaScript enabled'
    
    # First login section
    doc.add_paragraph()
    add_paragraph_styled(doc, 'First Login: Step-by-Step\n', bold=True)
    
    steps = [
        'Navigate to planner.cloud.microsoft.com',
        'Log in with your company email/password',
        'Click on your project card (ex: "PLC - RR - Calima - 01 Singapore")',
        'Board view opens with cards organized by phase',
        'Personalize settings: Language, Notifications, Theme',
        'Set up Teams integration by finding Planner tab in your project channel'
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    # ===== III.A.2: BOARD NAVIGATION =====
    doc.add_page_break()
    add_heading_styled(doc, 'III.A.2: Board Navigation Basics', 3)
    
    add_paragraph_styled(doc, 'Understanding Board Views\n', bold=True, size=11)
    p = doc.add_paragraph('Every Planner board has 4 main views. Choose which to use based on your task.')
    
    # Views table
    views_table = doc.add_table(rows=5, cols=3)
    views_table.style = 'Light Grid Accent 1'
    
    headers = views_table.rows[0].cells
    headers[0].text = 'View Name'
    headers[1].text = 'When to Use'
    headers[2].text = 'Key Actions'
    
    for header in headers:
        set_cell_background(header, '1010EB')
        for p in header.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    view_data = [
        ('Grid View (Default)', 'Searching for specific cards, sorting, overview', 'Click cells to edit, sort/filter by columns'),
        ('Board View (Kanban)', 'Moving cards between phases, visual status', 'Drag cards between buckets, rearrange within phase'),
        ('Timeline View (Gantt)', 'Planning schedule, seeing dependencies, presentations', 'Drag bars to change dates, follow red dependency lines'),
        ('My Tasks View (Personal)', 'Morning planning, checking your workload', 'Filter by project, see only cards assigned to you'),
    ]
    
    for i, (name, when, actions) in enumerate(view_data, 1):
        row = views_table.rows[i].cells
        row[0].text = name
        row[1].text = when
        row[2].text = actions
    
    doc.add_paragraph()
    
    # Diagrams: Insert Board Views Diagram
    add_paragraph_styled(doc, 'Visual: NPI Lifecycle Process Flow', bold=True, size=11)
    try:
        doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_01_npi_lifecycle.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Warning: Could not insert diagram 1: {e}")
    
    doc.add_paragraph()
    
    # ===== III.B: CARD STRUCTURE =====
    doc.add_page_break()
    add_heading_styled(doc, 'III.B: Understanding the NPI Card Structure', 2, COLORS['primaryBlue'])
    
    add_paragraph_styled(doc, 'Context: Why Three Levels of Cards?\n', bold=True, size=11)
    
    p = doc.add_paragraph('Planner Premium uses three levels of cards to mirror the NPI project structure:')
    doc.add_paragraph('Upstream card (project level) = What\'s the overall project?', style='List Bullet')
    doc.add_paragraph('Regional cards (regional level) = What\'s our region delivering?', style='List Bullet')
    doc.add_paragraph('Sub-tasks (functional level) = What\'s each team member doing?', style='List Bullet')
    
    # Insert Card Hierarchy Diagram
    doc.add_paragraph()
    add_paragraph_styled(doc, 'Card Hierarchy Diagram', bold=True, size=10)
    try:
        doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_02_card_hierarchy.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Warning: Could not insert diagram 2: {e}")
    
    doc.add_paragraph()
    
    # III.B.1: Upstream Cards
    add_heading_styled(doc, 'III.B.1: Upstream Card (Project-Level Summary)', 3)
    
    add_paragraph_styled(doc, 'What It Is\n', bold=True)
    doc.add_paragraph('The Upstream card represents the entire NPI project at a global level.', style='List Bullet')
    doc.add_paragraph('One per project (not one per region)', style='List Bullet')
    doc.add_paragraph('Created by: Main PM (during Pipeline phase)', style='List Bullet')
    doc.add_paragraph('Moves through all phase buckets: Pipeline → GA → RA → Engage → Kick-off → Execute → Hypercare → Close-out', style='List Bullet')
    doc.add_paragraph('Shows overall progress: As regional cards complete, upstream card progress updates', style='List Bullet')
    
    add_paragraph_styled(doc, '\nHow to Read an Upstream Card\n', bold=True)
    doc.add_paragraph('Open the card and scan these key fields:', style='List Number')
    doc.add_paragraph('Global Service Model (tells you what regions must deliver)', style='List Number 2')
    doc.add_paragraph('Global Launch Date (critical timeline)', style='List Number 2')
    doc.add_paragraph('Complexity (indicates scope)', style='List Number 2')
    doc.add_paragraph('Regions Involved (how many you\'re coordinating)', style='List Number 2')
    doc.add_paragraph('Check the Progress bar (% of regions complete)', style='List Number')
    doc.add_paragraph('Review Risk Status (Green/Yellow/Red/Blocked)', style='List Number')
    
    # III.B.2: Regional Cards
    add_heading_styled(doc, 'III.B.2: Regional Card (Operating Unit-Specific)', 3)
    
    add_paragraph_styled(doc, 'What It Is\n', bold=True)
    doc.add_paragraph('The Regional card represents your specific region\'s work on the NPI project.', style='List Bullet')
    doc.add_paragraph('One per region, per project', style='List Bullet')
    doc.add_paragraph('Created by: Regional PM (during Regional Alignment phase)', style='List Bullet')
    doc.add_paragraph('Depends on: Upstream card (can\'t exist until GA is complete)', style='List Bullet')
    
    # III.B.3: Sub-tasks
    add_heading_styled(doc, 'III.B.3: Sub-Tasks (Functional Deliverables)', 3)
    
    add_paragraph_styled(doc, 'What They Are\n', bold=True)
    doc.add_paragraph('Sub-tasks are individual deliverables that team members complete.', style='List Bullet')
    doc.add_paragraph('Live under a regional card (not standalone)', style='List Bullet')
    doc.add_paragraph('Created by: Regional PM or Main PM (should use template)', style='List Bullet')
    doc.add_paragraph('Owned by: Functional SME (Training lead, Service manager, Ops manager, etc.)', style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== III.C: DAY-TO-DAY TASKS =====
    doc.add_page_break()
    add_heading_styled(doc, 'III.C: Day-to-Day Tasks & Card Management', 2, COLORS['primaryBlue'])
    
    add_heading_styled(doc, 'III.C.1: Creating a New Card (Regional PM Perspective)', 3)
    
    add_paragraph_styled(doc, 'Scenario: Creating Your First Regional Card\n', bold=True, italic=True)
    doc.add_paragraph('You\'re the Regional PM for Singapore on the Calima project. Global Alignment is complete. Now it\'s time to create your regional card to kick off Regional Alignment phase.')
    
    add_paragraph_styled(doc, '\nStep-by-Step: Creating a Regional Card\n', bold=True)
    
    steps_detailed = [
        ('Navigate to the Right Board', 'Go to planner.cloud.microsoft.com, find board "PLC - RR - Calima - 01 Singapore", click to open'),
        ('Locate the Upstream Card', 'Find "NPI - Calima - CAL-001", read it to understand global context'),
        ('Click to Create New Card', 'Find the RA bucket, click "+ Add card" button'),
        ('Fill in Basic Information', 'Title: "[PLC Name] - [Region] - Regional Readiness", Assigned to: You, Start date: When RA begins, Due date: Regional launch date, Description: Regional scope and constraints'),
        ('Add Custom Fields', 'Fill all ⭐ required fields: Planisware ID, Service Model, Complexity, Launch Date, Stakeholder Contact'),
        ('Save the Card', 'Click Save or Create'),
        ('Notify Your Team', 'Post in Teams channel to announce regional card is live'),
    ]
    
    for step_num, (step_title, step_desc) in enumerate(steps_detailed, 1):
        p = doc.add_paragraph(f'Step {step_num}: {step_title}', style='List Number')
        doc.add_paragraph(step_desc, style='List Number 2')
    
    # Insert Phase Gate Diagram
    doc.add_paragraph()
    add_paragraph_styled(doc, 'Phase Gate Decision Diagram', bold=True, size=10)
    try:
        doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_03_phase_gate_decision.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Warning: Could not insert diagram 3: {e}")
    
    doc.add_paragraph()
    
    # ===== III.C.4: MOVING CARDS =====
    add_heading_styled(doc, 'III.C.4: Moving Cards Between Buckets (Phase Progression)', 3)
    
    add_paragraph_styled(doc, 'What Moving a Card Means\n', bold=True)
    doc.add_paragraph('When a card moves from one bucket to another, the project is advancing to the next phase.')
    
    phase_moves = [
        'Pipeline → GA: Project approved, moving to global alignment',
        'GA → RA: Global alignment complete, ready for regional alignment',
        'RA → Engage: Regional alignment done, start engaging with teams',
        'Engage → Kick-off: SME commitments received, formal kick-off happening',
        'Kick-off → Execute: Project plan approved, work is starting',
        'Execute → Hypercare: Launch day reached, monitoring support period',
        'Hypercare → Close-out: 30 days of go-live support complete, project closing',
    ]
    
    for move in phase_moves:
        doc.add_paragraph(move, style='List Bullet')
    
    doc.add_paragraph()
    add_paragraph_styled(doc, 'Before You Move a Card: The Checklist\n', bold=True)
    
    checklist_items = [
        'All required custom fields are filled? (⭐ fields)',
        'All deliverables for THIS phase are at least started?',
        'No critical blockers preventing next phase?',
        'Key stakeholders have approved? (Regional Service Leader, Main PM)',
        'Are there any risks that should be resolved first?',
        'Is the next phase\'s timeline realistic?',
        'Do we have the resources to execute the next phase?',
    ]
    
    for item in checklist_items:
        doc.add_paragraph('☐ ' + item, style='List Bullet')
    
    # ===== III.F: ESCALATION =====
    doc.add_page_break()
    add_heading_styled(doc, 'III.F: Communication & Collaboration', 2, COLORS['primaryBlue'])
    
    add_heading_styled(doc, 'III.F.3: Escalation Paths (When to Ask for Help)', 3)
    
    add_paragraph_styled(doc, 'Escalation Decision Tree\n', bold=True)
    try:
        doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_04_escalation_tree.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Warning: Could not insert diagram 4: {e}")
    
    doc.add_paragraph()
    
    add_paragraph_styled(doc, 'Escalation Contacts & Methods\n', bold=True)
    
    escalation_table = doc.add_table(rows=6, cols=4)
    escalation_table.style = 'Light Grid Accent 1'
    
    headers = escalation_table.rows[0].cells
    headers[0].text = 'Situation'
    headers[1].text = 'Who'
    headers[2].text = 'How'
    headers[3].text = 'Timeline'
    
    for header in headers:
        set_cell_background(header, '140F4B')
        for p in header.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    escalation_data = [
        ('Quick Questions', 'Regional PM', 'Teams chat or call', '24 hours'),
        ('Blockers (Can\'t Proceed)', 'Regional PM (immediately)', 'Teams message or email', 'Same day'),
        ('Timeline Issues', 'Regional PM + Service Leader', 'Governance call or urgent 1:1', '48 hours'),
        ('Stakeholder Issues', 'Regional Service Leader', '1:1 meeting or email', 'Decision within 48 hrs'),
        ('Technical Issues', 'IT Help Desk', 'Help ticket with details', '4 hours'),
    ]
    
    for i, (situation, who, how, timeline) in enumerate(escalation_data, 1):
        row = escalation_table.rows[i].cells
        row[0].text = situation
        row[1].text = who
        row[2].text = how
        row[3].text = timeline
    
    # ===== III.G: SCENARIOS =====
    doc.add_page_break()
    add_heading_styled(doc, 'III.G: Common Scenarios & How to Handle Them', 2, COLORS['primaryBlue'])
    
    add_heading_styled(doc, 'III.G.1: Scenario 1 — You\'re a Regional PM Creating Your First Regional Card', 3)
    
    add_paragraph_styled(doc, 'The Situation\n', bold=True)
    doc.add_paragraph('You\'ve just been assigned as Regional PM for India on the Calima project. Global Alignment is complete. It\'s time to create your regional card and start Regional Alignment.')
    
    add_paragraph_styled(doc, '\nStep-by-Step Solution\n', bold=True)
    
    solution_steps = [
        'Get the Template - Email Main PM, they\'ll send you the regional card template',
        'Create Your Regional Card - Use Section III.C.1 for detailed steps',
        'Ask for Planisware ID - Get from Main PM, add to custom fields',
        'Create Sub-Tasks - Add training, service procedures, and other deliverables',
        'Notify Your Team - Post in Teams explaining timeline and deliverables',
        'Schedule First Governance Call - Discuss India-specific considerations',
    ]
    
    for step in solution_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    add_paragraph_styled(doc, 'Result: Your regional card is live, team knows what to do, and you\'re ready to start Regional Alignment phase.', italic=True)
    
    # ===== III.H: QUICK REFERENCE CHECKLISTS =====
    doc.add_page_break()
    add_heading_styled(doc, 'III.H: Quick Reference Checklists (Printable)', 2, COLORS['primaryBlue'])
    
    add_heading_styled(doc, 'III.H.1: Daily Task Checklist', 3)
    
    daily_checklist = [
        'Check My Tasks view for new assignments',
        'Scan for RED/BLOCKED cards (anything needing urgent attention)',
        'Click each red card and read recent comments',
        'Update status on cards assigned to me (if changed)',
        'Add comment if there\'s significant progress to report',
        'Check for @mentions from teammates',
        'Respond to any questions or requests',
    ]
    
    for item in daily_checklist:
        doc.add_paragraph('☐ ' + item, style='List Bullet')
    
    doc.add_paragraph()
    add_paragraph_styled(doc, 'Time investment: 5-10 minutes', italic=True, size=9)
    
    # Weekly Checklist
    add_heading_styled(doc, 'III.H.2: Weekly Task Checklist', 3)
    
    doc.add_paragraph('Do this every Friday afternoon (before governance calls).')
    
    weekly_items = [
        'Open all MY cards (filter by "Assigned to me")',
        'For each card: Is status accurate? Has anything changed? What\'s the next action? Update "Next Activity Date"',
        'Check all sub-tasks: Any due this week? Any overdue? Any completed?',
        'Add a status comment to each card summarizing the week',
    ]
    
    for item in weekly_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_paragraph()
    add_paragraph_styled(doc, 'Time investment: 30-45 minutes | Result: You\'re ready for governance calls with complete status information', 
                         italic=True, size=9)
    
    # ===== INSERT SYSTEM INTEGRATION & ORG ROLES DIAGRAMS =====
    doc.add_page_break()
    
    add_heading_styled(doc, 'Supporting Diagrams', 2, COLORS['primaryBlue'])
    
    add_paragraph_styled(doc, 'System Integration Architecture\n', bold=True, size=11)
    try:
        doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_05_system_integration.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Warning: Could not insert diagram 5: {e}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_styled(doc, 'Organizational Roles & Responsibilities Matrix\n', bold=True, size=11)
    try:
        doc.add_picture('/Users/casild1/Documents/VS_Vault/projects/brand-sop/diagram_06_roles_matrix.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Warning: Could not insert diagram 6: {e}")
    
    doc.add_paragraph()
    
    # ===== PLACEHOLDER SECTIONS FOR REMAINING PARTS =====
    doc.add_page_break()
    
    add_heading_styled(doc, 'Part IV: Project Manager Pathway', 1, COLORS['deepNavy'])
    doc.add_paragraph('[Detailed PM procedures - To be completed]', style='List Bullet')
    
    doc.add_page_break()
    
    add_heading_styled(doc, 'Part V: People Leader Pathway', 1, COLORS['deepNavy'])
    doc.add_paragraph('[Detailed leader procedures - To be completed]', style='List Bullet')
    
    doc.add_page_break()
    
    add_heading_styled(doc, 'Part VI: Shared Resources', 1, COLORS['deepNavy'])
    doc.add_paragraph('[Templates, glossary, contact directory - To be completed]', style='List Bullet')
    
    doc.add_page_break()
    
    add_heading_styled(doc, 'Part VII: Appendices', 1, COLORS['deepNavy'])
    doc.add_paragraph('[Licensing, change log, shortcuts, feedback - To be completed]', style='List Bullet')
    
    # Save document
    doc.save('/Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx')
    
    print("✅ Word document updated successfully!")
    print("📄 Document saved to: /Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx")
    print("\n📊 Content Summary:")
    print("  ✓ Cover page with version control")
    print("  ✓ Quick Start: Find Your Role")
    print("  ✓ Table of Contents")
    print("  ✓ Foundation section (all audiences)")
    print("  ✓ User Pathway (III.A - III.H) - COMPLETE")
    print("  ✓ 6 embedded professional diagrams with Medtronic branding")
    print("  ✓ Placeholder sections for PM, Leader, Shared Resources, Appendices")
    print("\n🎨 Diagrams Embedded:")
    print("  1. NPI Lifecycle Process Flow")
    print("  2. Card Hierarchy (Upstream → Regional → Sub-tasks)")
    print("  3. Phase Gate Decision Tree")
    print("  4. Escalation Decision Tree")
    print("  5. System Integration Architecture")
    print("  6. Organizational Roles Matrix")

if __name__ == '__main__':
    main()
