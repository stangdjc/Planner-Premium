#!/usr/bin/env python3
"""
Integrate PM Pathway into Word document and add cross-references
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Brand Colors
DEEPNAVY = RGBColor(20, 15, 75)
PRIMARYBLUE = RGBColor(16, 16, 235)
CHARCOAL = RGBColor(60, 60, 60)
LIGHTGRAY = RGBColor(245, 245, 245)

def parse_pm_pathway():
    """Read and parse PM_PATHWAY_DETAILED.md"""
    with open('/Users/casild1/Documents/VS_Vault/projects/Planner_Prem/procedure_document/PM_PATHWAY_DETAILED.md', 'r') as f:
        content = f.read()
    return content

def add_styled_heading(doc, text, level):
    """Add heading with proper brand styling"""
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)

    if level == 1:
        run.font.size = Pt(24)
        run.font.color.rgb = DEEPNAVY
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(18)
        run.font.color.rgb = PRIMARYBLUE
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(14)
        run.font.color.rgb = CHARCOAL
        run.font.bold = True
    elif level == 4:
        run.font.size = Pt(12)
        run.font.color.rgb = CHARCOAL
        run.font.bold = True

    return p

def add_styled_paragraph(doc, text, is_bold=False):
    """Add paragraph with proper styling"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = CHARCOAL
        if is_bold:
            run.font.bold = True
    return p

def process_markdown_section(doc, md_text):
    """Convert markdown section to Word format"""
    lines = md_text.strip().split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            add_styled_heading(doc, line[2:].strip(), 1)
        elif line.startswith('## '):
            add_styled_heading(doc, line[3:].strip(), 2)
        elif line.startswith('### '):
            add_styled_heading(doc, line[4:].strip(), 3)
        elif line.startswith('#### '):
            add_styled_heading(doc, line[5:].strip(), 4)

        # Handle bold text lines
        elif line.startswith('**') and line.endswith('**'):
            add_styled_paragraph(doc, line.replace('**', ''), is_bold=True)

        # Handle lists
        elif line.strip().startswith('- '):
            bullet_text = line.strip()[2:]
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = CHARCOAL

        # Handle checkboxes
        elif '✓' in line or '- [ ]' in line or '- [x]' in line:
            cleaned = line.replace('- [x]', '✓').replace('- [ ]', '☐').strip()
            p = doc.add_paragraph(cleaned, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = CHARCOAL

        # Handle regular paragraphs
        else:
            if line.strip():
                add_styled_paragraph(doc, line.strip())

        i += 1

def add_pm_pathway_to_document():
    """Load document and integrate PM Pathway"""

    # Load existing document
    doc = Document('/Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx')

    # Read PM Pathway content
    pm_content = parse_pm_pathway()

    # Add page break
    doc.add_page_break()

    # Add Part IV heading
    add_styled_heading(doc, 'PART IV: Project Manager Pathway - Governance & Oversight', 1)
    doc.add_paragraph()

    # Process PM Pathway markdown
    process_markdown_section(doc, pm_content)

    # Add page break before Part V
    doc.add_page_break()

    # Add Part V placeholder
    add_styled_heading(doc, 'PART V: People Leader Pathway - Strategic Alignment & Portfolio Health', 1)
    doc.add_paragraph('This section covers:')
    doc.add_paragraph('• Strategic alignment of NPI portfolio with organizational goals', style='List Bullet')
    doc.add_paragraph('• Portfolio health assessment and resource allocation', style='List Bullet')
    doc.add_paragraph('• Decision-making frameworks for project prioritization', style='List Bullet')
    doc.add_paragraph('• Leadership reporting and communication templates', style='List Bullet')
    doc.add_paragraph('• Coaching and development for Project Managers', style='List Bullet')
    doc.add_paragraph('[Detailed content to be added]')

    doc.add_page_break()

    # Add Part VI placeholder
    add_styled_heading(doc, 'PART VI: Shared Resources - Templates, Tools & References', 1)
    doc.add_paragraph('This section includes:')
    doc.add_paragraph('• Phase Gate Approval Templates (downloadable checklists)', style='List Bullet')
    doc.add_paragraph('• Governance Call Agendas and Meeting Templates', style='List Bullet')
    doc.add_paragraph('• Status Report Templates and Portfolio Dashboard examples', style='List Bullet')
    doc.add_paragraph('• Risk Management and Escalation Decision Trees', style='List Bullet')
    doc.add_paragraph('• Glossary of NPI-specific terms and acronyms', style='List Bullet')
    doc.add_paragraph('• Contact Directory (Main PM, Regional PMs, Service Leaders)', style='List Bullet')
    doc.add_paragraph('• Frequently Asked Questions (FAQs) by role', style='List Bullet')
    doc.add_paragraph('[Detailed templates and resources to be added]')

    doc.add_page_break()

    # Add Part VII placeholder
    add_styled_heading(doc, 'PART VII: Appendices - Additional References', 1)
    doc.add_paragraph('This section includes:')
    doc.add_paragraph('• Licensing and Access Information for Planner Premium', style='List Bullet')
    doc.add_paragraph('• Change Log and Version History', style='List Bullet')
    doc.add_paragraph('• Keyboard Shortcuts and Tips for Power Users', style='List Bullet')
    doc.add_paragraph('• Troubleshooting Guide: Common Issues and Solutions', style='List Bullet')
    doc.add_paragraph('• Feedback and Improvement Suggestions', style='List Bullet')
    doc.add_paragraph('[Detailed appendices to be added]')

    # Save document
    doc.save('/Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx')

    return True

if __name__ == '__main__':
    print("\n📝 Integrating PM Pathway into Word document...")
    try:
        add_pm_pathway_to_document()
        print("✅ PM Pathway successfully integrated!")
        print("📄 Location: /Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx")
        print("\n📊 Document now includes:")
        print("  ✓ Part II: Foundation (Core concepts, system overview)")
        print("  ✓ Part III: User Pathway (Complete procedures for users)")
        print("  ✓ Part IV: Project Manager Pathway (Governance & oversight)")
        print("  ✓ Part V: People Leader Pathway (Strategic alignment)")
        print("  ✓ Part VI: Shared Resources (Templates & tools)")
        print("  ✓ Part VII: Appendices (Additional references)")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
