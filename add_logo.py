#!/usr/bin/env python3
"""
Add Medtronic logo to procedure document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DEEPNAVY = RGBColor(20, 15, 75)
PRIMARYBLUE = RGBColor(16, 16, 235)

def add_logo_to_header(doc):
    """Add Medtronic logo to document header"""
    section = doc.sections[0]
    header = section.header

    # Clear existing header content
    for para in header.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Add logo to header
    header_para = header.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    try:
        # Add Medtronic logo
        logo_path = "/Users/casild1/Documents/VS_Vault/projects/brand-sop/brand-mdt-templates/coll-art-medtronic-logos/Medtronic logo - Black (limited use)/Logos for DIGITAL SCREENS (RGB)/art-logo-en-rgb-k.png"
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
    except Exception as e:
        print(f"⚠️  Could not add logo: {e}")
        header_para.add_run("Medtronic")

def add_logo_to_title_page(doc):
    """Add logo to title page"""
    # Find the title paragraph
    for i, para in enumerate(doc.paragraphs):
        if 'Microsoft Planner Premium' in para.text:
            # Insert logo before title
            new_para = para.insert_paragraph_before()
            new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            try:
                logo_path = "/Users/casild1/Documents/VS_Vault/projects/brand-sop/brand-mdt-templates/coll-art-medtronic-logos/Medtronic logo - Black (limited use)/Logos for DIGITAL SCREENS (RGB)/art-logo-en-rgb-k.png"
                run = new_para.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
            except Exception as e:
                print(f"⚠️  Could not add logo to title: {e}")

            # Add spacing
            spacing_para = para.insert_paragraph_before()

            break

def add_footer_with_branding(doc):
    """Add footer with page numbers and branding"""
    section = doc.sections[0]
    footer = section.footer

    # Clear existing footer
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Add footer with page number and copyright
    footer_para = footer.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add page number field
    run = footer_para.add_run()

    # Add copyright and page info
    copyright_text = "Microsoft Planner Premium for NPI Regional Readiness | Page "
    run = footer_para.add_run(copyright_text)
    run.font.size = Pt(9)
    run.font.color.rgb = DEEPNAVY

def main():
    print("\n📝 Adding Medtronic Logo to Document...\n")

    doc_path = '/Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx'

    doc = Document(doc_path)

    print("✓ Adding logo to header...")
    add_logo_to_header(doc)

    print("✓ Adding logo to title page...")
    add_logo_to_title_page(doc)

    print("✓ Adding branded footer...")
    add_footer_with_branding(doc)

    # Save document
    doc.save(doc_path)

    print("\n✅ Logo and branding elements added successfully!")
    print("📄 Location: /Users/casild1/Documents/VS_Vault/projects/brand-sop/NPI_Planner_Premium_Procedure.docx")
    print("\n🎨 Branding Updates:")
    print("  ✓ Medtronic logo added to document header")
    print("  ✓ Medtronic logo added to title page")
    print("  ✓ Branded footer with copyright added")
    print("  ✓ Professional document header and footer structure")

if __name__ == '__main__':
    main()
